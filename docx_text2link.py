#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# docx_text2link.py
#   Convert specific text in a DOCX document to hyperlinks,
#   without changing anything else.
#
# by M.H.V. Werts, May 2020
# using code snippets by others cited below
#
# USE AT YOUR OWN RISK!
# This program has only been very partially tested
#
# A typical use case would be that you have a document
# with a formatted bibliography (e.g. using Zotero followed
# by *unlinking* the references) and you need to supply
# clickable hyperlinks to the corresponding papers on-line 
# (e.g. using their Digital Object Identifiers (DOI).
#
# Requires python-docx (we used 0.8.10)
#
# In conda-forge channel
#   conda install python-docx
#
# or (elsewhere)
#   pip3 install python-docx
#
#
# Writing of this script was possible thanks to the following information:
# [1]  https://github.com/python-openxml/python-docx/issues/74
# [2]  https://stackoverflow.com/questions/40475757/how-to-extract-the-url-in-hyperlinks-from-a-docx-file-using-python
# [3]  https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
#
# The following very helpful code was used in the script:
# [4]  https://github.com/python-openxml/python-docx/issues/519#issuecomment-441710870
# [5]  https://github.com/python-openxml/python-docx/issues/74#issuecomment-441351994
#
# Method [3] provides alternative way of inserting links. However, this 
# generates some complications with formatting, and links can only
# appear at end of paragraph.


import sys
import re
import docx
from docx.shared import RGBColor

#############################################

# get command line parameters

if len(sys.argv)!=3:
    print('usage: '+sys.argv[0]+' <input file> <output file>')
    exit(1);

fpname = sys.argv[1]
outfpn = sys.argv[2]

#############################################

# configurable parameters (#TODO: configuration file/CL options)
# refmarker = '\[\d*\].*'  # regexp to recognize a bibliography item: "[decimal number]"
refmarker = '\[\d.*'  # regexp to recognize a bibliography item: "[decimal"
doistr = 'DOI:' #pre-fix to be recognised for link creation
urlbase = 'https://dx.doi.org/' #pre-fix for assembling link
linkrgb = '0563C1' #colour of links
linkunderline = True #underline links?

# parameter-derived variables
linkRGBcolor = RGBColor.from_string(linkrgb)

#############################################

def add_run_copy(paragraph, run, text=None):
    """Add a 'run' at the end of a paragraph using a 'run' template.
    
    A 'run' is a part of a paragraph with identically formatted text
    each time character formatting changes, a new 'run' is needed.
    
    source:
    https://github.com/python-openxml/python-docx/issues/519#issuecomment-441710870
    
    I have the impression that not everything is neatly copied
    Don't know why
    
    """
    r = paragraph.add_run(text=run.text if text is None else text, style=run.style)
    r.bold = run.bold
    r.italic = run.italic
    r.underline = run.underline
    r.font.all_caps = run.font.all_caps
    r.font.bold = run.font.bold
    r.font.color.rgb = run.font.color.rgb
    r.font.color.theme_color = run.font.color.theme_color
    #r.font.color.type = run.font.color.type
    r.font.complex_script = run.font.complex_script
    r.font.cs_bold = run.font.cs_bold
    r.font.cs_italic = run.font.cs_italic
    r.font.double_strike = run.font.double_strike
    r.font.emboss = run.font.emboss
    r.font.hidden = run.font.hidden
    r.font.highlight_color = run.font.highlight_color
    r.font.imprint = run.font.imprint
    r.font.italic = run.font.italic
    r.font.math = run.font.math
    r.font.name = run.font.name
    r.font.no_proof = run.font.no_proof
    r.font.outline = run.font.outline
    r.font.rtl = run.font.rtl
    r.font.shadow = run.font.shadow
    r.font.size = run.font.size
    r.font.small_caps = run.font.small_caps
    r.font.snap_to_grid = run.font.snap_to_grid
    r.font.spec_vanish = run.font.spec_vanish
    r.font.strike = run.font.strike
    r.font.subscript = run.font.subscript
    r.font.superscript = run.font.superscript
    r.font.underline = run.font.underline
    r.font.web_hidden = run.font.web_hidden
    return r


def add_hyperlink_into_run(paragraph, run, url):
    """ Convert a 'run' inside a paragraph to a hyperlink.
    
    does not do formatting, preserves style
    
    source:
    https://github.com/python-openxml/python-docx/issues/74#issuecomment-44135199

    """
    runs = paragraph.runs
    for i in range(len(runs)):
        if runs[i].text == run.text:
            break
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    hyperlink.append(run._r)
    paragraph._p.insert(i+1,hyperlink)


# GO!
d = docx.Document(fpname)                  

# look for paragraphs containing references with DOI
# the search expression may need to be changed for different use cases
# using regexp (import re)
# use match to match at start of text
rex = re.compile(refmarker+doistr)
psels = []
for p in d.paragraphs:
    txt = p.text
    if rex.match(txt):
        psels.append(p)

# inspect selected paragraphs before processing
print('=================')
print('BEFORE PROCESSING')
print('=================')
for p in psels:
    print(p.text)


print('')
print('==========')
print('PROCESSING')
print('==========')

# now in each selected paragraph, apply magic
for p in psels:
    irdoi = 0
    rdoifound = False
    # The following loop will bug if doistr is not present!
    # (but it should be present after correct selection)
    #
    # If no doistr found, then the reference is just skipped
    #
    # Unhandled case: 
    # When the 'doistr' is distributed over several 'runs',
    # then it cannot be found using the present algorithm.
    # This may happen because of Word autoformatting.
    # For now, such errors should be prevented by correcting
    # the formatting within the document (select, cut, and paste
    # as plain text) 
    #TODO: use exception handling, report and ignore any offending
    # occurrences
    while (not rdoifound) and (irdoi<len(p.runs)):
        r = p.runs[irdoi]
        ixdoi = r.text.find(doistr)
        if (ixdoi>=0):
            rdoifound = True
        else:
            irdoi += 1
    if rdoifound:
        r = p.runs[irdoi]
        txt1 = r.text[:ixdoi]+'('
        txt2 = r.text[ixdoi+4:].strip()
        rnew=add_run_copy(p, r, text='link')
        rend=add_run_copy(p, r, text=')')
        # tune colour / underline
        # this is a bit strange, but it works
        rnew.font.color.rgb = linkRGBcolor
        rnew.font.underline = linkunderline
        url = urlbase+txt2
        add_hyperlink_into_run(p, rnew, url)
        r.text = txt1
    else:
        print('skipping 1 DOI reference...')

# inspect selected paragraphs before processing
print('')
print('================')
print('AFTER PROCESSING')
print('================')
for p in psels:
    print(p.text)           

d.save(outfpn)

